import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC=Path('OpenEvidence_Competitive_Intelligence_Report_2026-07-25.md')
OUT=Path('OpenEvidence_Competitive_Intelligence_Report_2026-07-25.docx')
D=Path('diagrams')
NAVY='153B5B'; BLUE='2E75B6'; GREEN='D9EAD3'; YELLOW='FFF2CC'; RED='F4CCCC'; LIGHT='D9EAF7'; GREY='F3F6F8'; TEXT='172B3A'

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def set_cell_margins(cell, top=80,start=80,bottom=80,end=80):
    tc = cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')
def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)
def set_cell_width(cell,width_inches):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(int(width_inches*1440))); tcW.set(qn('w:type'),'dxa')
def add_hyperlink(paragraph, text, url, color='0563C1'):
    part=paragraph.part
    r_id=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    hyperlink=OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'),r_id)
    new_run=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    c=OxmlElement('w:color'); c.set(qn('w:val'),color); rPr.append(c)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    new_run.append(rPr); t=OxmlElement('w:t'); t.text=text; new_run.append(t); hyperlink.append(new_run); paragraph._p.append(hyperlink)
def add_runs_md(p,text,font_size=None):
    # handles simple markdown link and bold, retains status emoji.
    pos=0
    token=re.compile(r'(\*\*.*?\*\*|\[.*?\]\(https?://[^)]+\))')
    for m in token.finditer(text):
        if m.start()>pos:
            r=p.add_run(text[pos:m.start()]);
            if font_size: r.font.size=Pt(font_size)
        t=m.group(0)
        if t.startswith('**'):
            r=p.add_run(t[2:-2]); r.bold=True
            if font_size: r.font.size=Pt(font_size)
        else:
            mm=re.match(r'\[(.*?)\]\((.*?)\)',t)
            add_hyperlink(p,mm.group(1),mm.group(2))
        pos=m.end()
    if pos<len(text):
        r=p.add_run(text[pos:]);
        if font_size: r.font.size=Pt(font_size)

def parse_table(lines):
    rows=[]
    for line in lines:
        if not line.strip().startswith('|'): continue
        cells=[x.strip() for x in line.strip().strip('|').split('|')]
        if all(re.fullmatch(r'[-: ]+',c or '-') for c in cells): continue
        rows.append(cells)
    return rows

def status_color(text):
    if '🟢' in text: return GREEN
    if '🟡' in text: return YELLOW
    if '🔴' in text: return RED
    return None

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.55); sec.right_margin=Inches(.55)
# styles
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos'); styles['Normal'].font.size=Pt(9.5); styles['Normal'].font.color.rgb=RGBColor.from_string(TEXT)
for level,size,color in [(1,18,NAVY),(2,14,BLUE),(3,11,NAVY)]:
    st=styles[f'Heading {level}']; st.font.name='Aptos Display'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos Display'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(14 if level==1 else 10); st.paragraph_format.space_after=Pt(5)
if 'Code Block' not in styles:
    st=styles.add_style('Code Block',WD_STYLE_TYPE.PARAGRAPH); st.font.name='DejaVu Sans Mono'; st.font.size=Pt(7.4); st.font.color.rgb=RGBColor.from_string(TEXT); st.paragraph_format.space_after=Pt(0); st.paragraph_format.space_before=Pt(0)
# header/footer
header=sec.header.paragraphs[0]; header.text='OpenEvidence — Competitive Intelligence & Ovexis Strategy Memo | Public-source snapshot: 25 July 2026'; header.style='Caption'; header.runs[0].font.color.rgb=RGBColor.from_string(BLUE)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('Confidential strategy research • Public sources only • ')
field=OxmlElement('w:fldSimple'); field.set(qn('w:instr'),'PAGE'); footer._p.append(field)
# Cover page
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
r=p.add_run('OPEN EVIDENCE'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Competitive Intelligence & Ovexis Strategy Memo'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Board-level public-source investigation'); r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(TEXT)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(26)
p.add_run('Target: OpenEvidence | Category: Clinical AI Assistant\nOfficial website: https://www.openevidence.com/\nResearch snapshot: 25 July 2026 (Asia/Kolkata)').font.size=Pt(11)
t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,LIGHT); set_cell_margins(c,180,180,180,180)
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=p.add_run('Evidence protocol\n'); rr.bold=True; rr.font.color.rgb=RGBColor.from_string(NAVY); rr.font.size=Pt(12)
p.add_run('🟢 Confirmed — direct public support\n🟡 Strong Inference — explicit, testable synthesis\n🔴 Speculation — unconfirmed scenario\n\nNo credentials were represented; no authenticated product areas, APIs or unauthorised systems were accessed.').font.size=Pt(10)
doc.add_paragraph('Prepared for Ovexis strategy, product, clinical governance and board discussion.\nCompanion workbook: OpenEvidence_Master_Feature_Inventory_2026-07-25.xlsx',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()
# Board diagrams section
h=doc.add_heading('Board diagrams',level=1)
p=doc.add_paragraph('🟢 These diagrams visualise the same evidence-labelled analysis in the report. Green nodes are publicly confirmed claims; yellow nodes are explicit inferences or Ovexis recommendations.'); p.paragraph_format.space_after=Pt(8)
diagrams=[
('Product Architecture Diagram','01_product_architecture.png'),
('AI Architecture Diagram','02_ai_rag_architecture.png'),
('Healthcare Data Flow Diagram','03_healthcare_data_flow.png'),
('User Journey Diagram','04_user_journey.png'),
('Feature Dependency Graph','05_feature_dependency_graph.png'),
('Business Model Canvas','06_business_model_canvas.png'),
]
for i,(cap,file) in enumerate(diagrams):
    doc.add_heading(cap,level=2)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(D/file),width=Inches(7.0))
    if i in (1,3): doc.add_page_break()
doc.add_page_break()

# Parse markdown
lines=SRC.read_text().splitlines()
i=0
in_code=False
code_lines=[]
while i<len(lines):
    line=lines[i]
    # Code block
    if line.startswith('```'):
        if not in_code:
            in_code=True; code_lines=[]
        else:
            p=doc.add_paragraph(style='Code Block')
            # preserve in a shaded single cell for clean display
            tb=doc.add_table(rows=1,cols=1); tb.alignment=WD_TABLE_ALIGNMENT.CENTER
            cell=tb.cell(0,0); shade(cell,GREY); set_cell_margins(cell,120,120,120,120)
            cp=cell.paragraphs[0]; cp.style='Code Block'; cp.add_run('\n'.join(code_lines))
            in_code=False
        i+=1; continue
    if in_code:
        code_lines.append(line); i+=1; continue
    # table block
    if line.strip().startswith('|'):
        block=[]
        while i<len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i]); i+=1
        rows=parse_table(block)
        if rows:
            n=max(len(r) for r in rows)
            tbl=doc.add_table(rows=0,cols=n); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit=False
            usable=7.35
            for ridx,row in enumerate(rows):
                cells=tbl.add_row().cells
                for j in range(n):
                    cell=cells[j]; set_cell_width(cell,usable/n); set_cell_margins(cell,60,65,60,65); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
                    value=row[j] if j<len(row) else ''
                    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
                    add_runs_md(p,value,7.2 if n>4 else 8.0)
                    if ridx==0:
                        shade(cell,NAVY)
                        for run in p.runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
                    else:
                        col=status_color(value)
                        if col: shade(cell,col)
                if ridx==0: set_repeat_table_header(tbl.rows[0])
            doc.add_paragraph().paragraph_format.space_after=Pt(1)
        continue
    # headings
    m=re.match(r'^(#{1,3})\s+(.*)$',line)
    if m:
        level=len(m.group(1)); txt=m.group(2)
        # avoid duplicate title after cover
        if not (level==1 and txt.startswith('OpenEvidence — Competitive Intelligence')):
            doc.add_heading(txt,level=level)
        i+=1; continue
    # horizontal rule
    if line.strip()=='---':
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:color'),'D9E2F3'); pbdr.append(bottom); pPr.append(pbdr)
        i+=1; continue
    # blank
    if not line.strip():
        i+=1; continue
    # bullets/numbered
    m=re.match(r'^(\d+\.)\s+(.*)$',line)
    if m:
        p=doc.add_paragraph(style='List Number'); add_runs_md(p,m.group(2)); i+=1; continue
    m=re.match(r'^[-*]\s+(.*)$',line)
    if m:
        p=doc.add_paragraph(style='List Bullet'); add_runs_md(p,m.group(1)); i+=1; continue
    # normal paragraph
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.07
    add_runs_md(p,line)
    i+=1

# Ensure all table fonts reasonable and header rows repeat
for tbl in doc.tables:
    for row_idx,row in enumerate(tbl.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs:
                    if r.font.size is None: r.font.size=Pt(8)
# add closing page

doc.add_page_break(); doc.add_heading('Deliverable control',level=1)
p=doc.add_paragraph('🟢 Main report: OpenEvidence_Competitive_Intelligence_Report_2026-07-25.docx and the equivalent Markdown file.\n🟢 Structured workbook: OpenEvidence_Master_Feature_Inventory_2026-07-25.xlsx with Master Feature Inventory, Evidence Register, Claim Ledger, Decision Ledger, Risk Register, Competitor Matrix, Roadmap Reconstruction and Business Model Canvas.\n🟢 Diagrams: companion PNG files in the diagrams folder.\n🟢 Research scope: public information through 25 July 2026; source boundaries and confidence are stated in the report.')
p.paragraph_format.space_after=Pt(8)

doc.core_properties.title='OpenEvidence Competitive Intelligence & Ovexis Strategy Memo'
doc.core_properties.subject='Public-source board strategy research'
doc.core_properties.author='Arena.ai Agent Mode'
doc.core_properties.comments='Public-source research snapshot, 25 July 2026.'
doc.save(OUT)
print('Wrote',OUT)
